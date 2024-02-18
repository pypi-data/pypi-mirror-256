import pandas as pd, numpy as np, seaborn as sns, matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import io
import warnings
warnings.filterwarnings('ignore')
from tqdm import tqdm

def null_Analysis(data_df, null_Cuttoff = 25, headcount = 5):
    print('Total Data Size {0} and total number of columns are {1}'.format((len(data_df)),(len(data_df.columns))))
    null_df = pd.DataFrame(columns=['Feature', 'Total Non-Null Value', 'No. of Null Value', '% NUll Value', 'Data Type'])
    feature = []
    total_val = []
    null_val = []
    per_null_val = []
    data_type = []
    for i in data_df.columns:
        feature.append(i)
        total_val.append(data_df[i].count())
        null_val.append(data_df[i].isnull().sum())
        per_null_val.append((data_df[i].isnull().sum() / len(data_df))*100)
        data_type.append(data_df[i].dtypes)
    null_df['Feature'] = feature
    null_df['Total  Non-Null Value'] = total_val
    null_df['No. of Null Value'] = null_val
    null_df['% NUll Value'] = per_null_val
    null_df['Data Type'] = data_type
    null_df =  null_df.sort_values(by='% NUll Value', axis=0, ascending=False)
    null_df = null_df.reset_index()
    null_df.drop('index', axis=1, inplace=True)
    high_null_column = []
    low_null_val = []
    for i,j in zip (null_df['Feature'], null_df['% NUll Value']):
        if j > null_Cuttoff:
            high_null_column.append(i)
        elif j < null_Cuttoff and j > 0:
            low_null_val.append(i)
    print('Columns with more than {0}% of null value : {1}'.format(null_Cuttoff, high_null_column) )
    print("")
    print('Columns with less than {0}% of null value : {1}'.format(null_Cuttoff, low_null_val) )
    print("")
    print("*************************")
    print("*** Dataframe Summary ***")
    print("*************************")
    for i in range (0, len(null_df)):
        print("Feature : {0} || Data Type : {1} ". format( null_df['Feature'].iloc[i], null_df['Data Type'].iloc[i]))
        print("No. of Null Value : {0} || Total Non-Null Value : {1} || % NUll Value : {2}".format(null_df['No. of Null Value'].iloc[i],null_df['Total  Non-Null Value'].iloc[i],null_df['% NUll Value'].iloc[i]))
        print('===========')


def plot_categorical_frequency(df, output_path='cat_graph.docx'):
    # Filter columns with categorical data
    categorical_columns = df.select_dtypes(include='object').columns
    document = Document()

    for column in categorical_columns:
        # Count the frequency of each category
        category_counts = df[column].value_counts()

        # If there is only one category, print the total count and skip plotting
        if len(category_counts) == 1:
            print(f"{column} has only one category: {category_counts.index[0]} (Total count: {category_counts.iloc[0]})")
            continue

        # If there are more than 10 categories, combine the rest into 'Others'
        if len(category_counts) > 10:
            top_categories = category_counts.head(10)
            other_categories_count = category_counts[10:].sum()
            category_counts = top_categories.append(pd.Series({'Others': other_categories_count}))

        # Plot the bar chart
        plt.figure(figsize=(10, 8))  # Adjusted figure size
        plt.subplots_adjust(bottom=0.2)  # Set bottom margin

        category_counts.plot(kind='bar', color='skyblue')
        plt.title(f'Frequency Distribution of {column}')
        plt.xlabel(column)
        plt.ylabel('Frequency')

        # Display the counts on top of the bars
        for i, count in enumerate(category_counts):
            plt.text(i, count + 0.1, str(count), ha='center')    

        buffer = io.BytesIO()
        plt.savefig(buffer, format='jpg')
        plt.close()

        document.add_picture(buffer, width=Inches(5), height=Inches(4))  # Adjusted image size

    document.save(output_path)


def convert_columns_to_preferred_types(df):
    for column in df.columns:
        original_dtype = df[column].dtype
        if pd.api.types.is_string_dtype(df[column]):
            try:
                df[column] = pd.to_datetime(df[column])
                print(f"Converted {column} to datetime")
            except ValueError:
                pass
        if pd.api.types.is_object_dtype(df[column]):
            try:
                df[column] = pd.to_numeric(df[column])
                print(f"Converted {column} to numeric")
            except ValueError:
                pass  
        new_dtype = df[column].dtype
        if original_dtype != new_dtype:
            print(f"{column} data type changed from {original_dtype} to {new_dtype}")
    return df


def generate_visualizations_and_report(df, output_path='output.docx'):
    """
    Generate visualizations and a statistical report for the DataFrame.

    Parameters:
    - df: pandas DataFrame
    - output_path: str, optional, default: 'output.docx'
        Path to save the output Word document.

    Returns:
    - None
    """

    document = Document()

    # Filter columns with numerical data
    numerical_columns = df.select_dtypes(include=['number']).columns
    categorical_columns = df.select_dtypes(include=['object']).columns

    stats_data = []  # List to accumulate statistical data for numerical columns
    with tqdm(total=len(df.columns), desc="Generating report") as pbar:
        for column in df.columns:
            pbar.set_postfix_str(f"Working on {column}")
            # Check if the column has only one category
            if column in categorical_columns:
                pbar.set_postfix_str(f"{column} is a categorical column")
                top_categories = df[column].value_counts().nlargest(10)
                other_categories_count = df[column].nunique() - top_categories.shape[0]
                if other_categories_count > 0:
                    top_categories['Others'] = other_categories_count
                plt.figure(figsize=(8, 6))  # Adjust the width of the plot
                ax = sns.barplot(x=top_categories.index, y=top_categories.values)
                plt.title(f'Count Plot for {column}')
                plt.xlabel(column)
                plt.ylabel('Count')
                plt.xticks(rotation=90)
                for p in ax.patches:
                    ax.annotate(format(p.get_height(), '.0f'),
                                (p.get_x() + p.get_width() / 2., p.get_height()),
                                ha='center', va='center',
                                xytext=(0, 5),
                                textcoords='offset points')
                plt.tight_layout()
                buffer = io.BytesIO()
                plt.savefig(buffer, format='jpg', bbox_inches='tight')  # Use bbox_inches='tight'
                plt.close()
                document.add_picture(buffer, width=Inches(5))
                document.add_paragraph("Figure: Count Plot for " + column)
                plt.show()

            elif column in numerical_columns:
                pbar.set_postfix_str(f"{column} is a numerical column")
                if df[column].nunique() > 1:  # Check if there are more than one unique value
                    # Line Plot
                    plt.figure(figsize=(8, 6))
                    plt.plot(df.index, df[column])
                    plt.title(f'Line Plot for {column}')
                    plt.xlabel('Index')
                    plt.ylabel(column)
                    plt.tight_layout()
                    buffer = io.BytesIO()
                    plt.savefig(buffer, format='jpg', bbox_inches='tight')
                    plt.close()
                    document.add_picture(buffer, width=Inches(5))
                    document.add_paragraph("Figure: Line Plot for " + column)
                    plt.show()

                    # Scatter Plot
                    plt.figure(figsize=(8, 6))
                    plt.scatter(df.index, df[column])
                    plt.title(f'Scatter Plot for {column}')
                    plt.xlabel('Index')
                    plt.ylabel(column)
                    plt.tight_layout()
                    buffer = io.BytesIO()
                    plt.savefig(buffer, format='jpg', bbox_inches='tight')
                    plt.close()
                    document.add_picture(buffer, width=Inches(5))
                    document.add_paragraph("Figure: Scatter Plot for " + column)
                    plt.show()

                    # Box Plot
                    plt.figure(figsize=(8, 6))
                    sns.boxplot(x=df[column])
                    plt.title(f'Box Plot for {column}')
                    plt.xlabel(column)
                    plt.ylabel('Values')
                    plt.tight_layout()
                    buffer = io.BytesIO()
                    plt.savefig(buffer, format='jpg', bbox_inches='tight')
                    plt.close()
                    document.add_picture(buffer, width=Inches(5))
                    document.add_paragraph("Figure: Box Plot for " + column)
                    plt.show()

                    try:
                        Q1 = df[column].quantile(0.25)
                        Q3 = df[column].quantile(0.75)
                        IQR = Q3 - Q1

                        stats_data.append([column, df[column].min(), df[column].max(), df[column].median(),
                                        df[column].std(), IQR])

                        # Calculate outliers
                        lower_bound = Q1 - 1.5 * IQR
                        upper_bound = Q3 + 1.5 * IQR
                        outliers_count = df[(df[column] < lower_bound) | (df[column] > upper_bound)].shape[0]
                        outliers_info = f"Number of outliers: {outliers_count} (lower bound: {lower_bound}, upper bound: {upper_bound})"
                        document.add_paragraph(outliers_info)

                    except Exception as e:
                        print(f"IQR calculation not possible for {column}. Skipping statistical summary.")
                        print(e)
                else:
                    print(f"Skipping statistical summary for {column} as it has only one unique value.")

            pbar.update(1)

    # Add a table to the document to display statistical data for numerical columns
    if stats_data:
        print("Adding statistical summary table...")
        document.add_heading('Statistical Summary', level=1)
        stats_table = document.add_table(rows=1, cols=7)  # Increased to 7 columns for outliers
        stats_table.style = 'Table Grid'
        hdr_cells = stats_table.rows[0].cells
        hdr_cells[0].text = 'Column Name'
        hdr_cells[1].text = 'Min'
        hdr_cells[2].text = 'Max'
        hdr_cells[3].text = 'Median'
        hdr_cells[4].text = 'Std Dev'
        hdr_cells[5].text = 'IQR'
        hdr_cells[6].text = 'Outliers'

        for stat in stats_data:
            row_cells = stats_table.add_row().cells
            row_cells[0].text = str(stat[0])
            row_cells[1].text = str(stat[1])
            row_cells[2].text = str(stat[2])
            row_cells[3].text = str(stat[3])
            row_cells[4].text = str(stat[4])
            row_cells[5].text = str(stat[5])
            # Number of outliers
            Q1 = stat[3]
            IQR = stat[5]
            lower_bound = Q1 - 1.5 * IQR
            upper_bound = Q1 + 1.5 * IQR
            outliers_count = df[(df[stat[0]] < lower_bound) | (df[stat[0]] > upper_bound)].shape[0]
            row_cells[6].text = str(outliers_count)

    document.save(output_path)